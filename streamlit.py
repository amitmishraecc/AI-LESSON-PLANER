import os
import bcrypt
import pymongo
import streamlit as st
from datetime import datetime
import json
import io
import base64

from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser

# Import from new structure
try:
    from src.config.settings import Settings
    from src.utils.export import generate_pdf, generate_word_doc, DOCX_AVAILABLE, REPORTLAB_AVAILABLE
    from src.utils.llm import LLM_Setup, generate_notes_and_quiz
    USE_MODULAR_STRUCTURE = True
except ImportError:
    # Fallback to inline functions if modules not found
    USE_MODULAR_STRUCTURE = False
    Settings = None

# Export libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import RGBColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    RGBColor = None


load_dotenv()

# --- Streamlit Config ---
st.set_page_config(
    page_title="AI Lesson Planner",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- MongoDB Connection ---
mongodb_uri = os.getenv('MONGODB_URI')

try:
    client = pymongo.MongoClient(mongodb_uri, serverSelectionTimeoutMS=10000)
    client.server_info()
    db = client["StudentDB"]
    users = db["users"]
    lesson_plans = db["lesson_plans"]  # Collection for saving lesson plans
except pymongo.errors.ServerSelectionTimeoutError:
    st.error("❌ Cannot connect to MongoDB.")
    if mongodb_uri == 'mongodb://localhost:27017/':
        st.info("💡 **For Local MongoDB:**\n1. Install MongoDB from https://www.mongodb.com/try/download/community\n2. Start MongoDB service\n3. Or use MongoDB Atlas (see below)")
    else:
        st.info("💡 **For MongoDB Atlas:**\n1. Check your connection string in the `.env` file (MONGODB_URI)\n2. Make sure your IP is whitelisted in Atlas\n3. Verify your username and password are correct")
    st.info("📝 **Setup MongoDB Atlas:**\n1. Go to https://www.mongodb.com/cloud/atlas\n2. Create a free cluster\n3. Get your connection string from 'Connect' → 'Connect your application'\n4. Add it to `.env` as: `MONGODB_URI=your_connection_string`")
    st.stop()
except pymongo.errors.ConfigurationError as e:
    st.error(f"❌ MongoDB connection string error: {str(e)}")
    st.info("💡 Check your MongoDB Atlas connection string format in the `.env` file")
    st.stop()
except pymongo.errors.OperationFailure as e:
    error_msg = str(e)
    if "Authentication failed" in error_msg or "bad auth" in error_msg.lower():
        st.error("❌ MongoDB Authentication Failed")
        st.warning("**Common causes:**")
        st.write("1. ❌ **Wrong username or password** - Verify in MongoDB Atlas → Database Access")
        st.write("2. 🔐 **Password needs URL encoding** - If password has special characters (@, #, $, etc.), they need to be encoded")
        st.write("3. 👤 **User doesn't exist** - Check Database Access in MongoDB Atlas")
        st.info("**How to fix:**\n1. Go to MongoDB Atlas → Database Access\n2. Verify your database user exists\n3. Reset password if needed\n4. If password has special characters, URL encode them:\n   - `@` becomes `%40`\n   - `#` becomes `%23`\n   - `$` becomes `%24`\n   - `%` becomes `%25`")
    else:
        st.error(f"❌ MongoDB operation error: {error_msg}")
    st.stop()
except Exception as e:
    error_msg = str(e)
    if "Authentication failed" in error_msg or "bad auth" in error_msg.lower():
        st.error("❌ MongoDB Authentication Failed")
        st.warning("**Please verify:**")
        st.write("1. Username and password in your connection string")
        st.write("2. Database user exists in MongoDB Atlas → Database Access")
        st.write("3. Password is correct (reset if needed)")
        st.info("💡 **Tip:** If your password contains special characters, you may need to URL encode them in the connection string")
    else:
        st.error(f"❌ MongoDB connection error: {error_msg}")
    st.stop()

# --- LLM Setup (Fallback if modular import fails) ---
if not USE_MODULAR_STRUCTURE:
    def LLM_Setup(prompt):
        api_key = os.getenv('key')
        if not api_key or api_key == 'your_groq_api_key_here':
            st.error("❌ Groq API key not found. Please set your API key in the .env file.")
            st.info("💡 **How to fix:**\n1. Create a `.env` file in the project root\n2. Add: `key=your_actual_groq_api_key`\n3. Get your API key from: https://console.groq.com/")
            st.stop()
        
        model = ChatGroq(
            model="llama-3.3-70b-versatile",
            groq_api_key=api_key,
            temperature=0.7
        )
        parser = StrOutputParser()
        output = model | parser
        output = output.invoke(prompt)
        return output

# --- Export Functions (Fallback if modular import fails) ---
if not USE_MODULAR_STRUCTURE:
    def generate_pdf(plan_data):
        """Generate PDF from lesson plan"""
        if not REPORTLAB_AVAILABLE:
            return None
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(0, 0, 139),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        title = Paragraph(f"{plan_data['subject']} - {plan_data['topic']}", title_style)
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=11,
            textColor=RGBColor(100, 100, 100),
            alignment=TA_CENTER
        )
        meta_text = f"<b>Grade/Level:</b> {plan_data['grade']} | <b>Duration:</b> {plan_data['duration']} | <b>Created:</b> {plan_data.get('created_at', 'N/A')}"
        story.append(Paragraph(meta_text, meta_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Content - Simple markdown to text conversion
        content = plan_data['content']
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                if line.startswith('# '):
                    style = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=18, spaceAfter=12, spaceBefore=12)
                    story.append(Paragraph(line[2:].strip(), style))
                elif line.startswith('## '):
                    style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, spaceAfter=10, spaceBefore=10)
                    story.append(Paragraph(line[3:].strip(), style))
                elif line.startswith('### '):
                    style = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=12, spaceAfter=8, spaceBefore=8)
                    story.append(Paragraph(line[4:].strip(), style))
                elif line.startswith('- ') or line.startswith('* '):
                    story.append(Paragraph(f"• {line[2:].strip()}", styles['Normal']))
                elif line.startswith('**') and line.endswith('**'):
                    bold_text = line.strip('*')
                    story.append(Paragraph(f"<b>{bold_text}</b>", styles['Normal']))
                else:
                    # Clean up markdown links
                    import re
                    line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_word_doc(plan_data):
        """Generate Word document from lesson plan"""
        if not DOCX_AVAILABLE:
            return None
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f"{plan_data['subject']} - {plan_data['topic']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Grade/Level: {plan_data['grade']}").bold = True
        meta_para.add_run(f" | Duration: {plan_data['duration']}").bold = True
        meta_para.add_run(f" | Created: {plan_data.get('created_at', 'N/A')}").bold = True
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Content
        content = plan_data['content']
        lines = content.split('\n')
        
        for line in lines:
            if line.strip():
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                elif line.startswith('**') and line.endswith('**'):
                    para = doc.add_paragraph()
                    para.add_run(line.strip('*')).bold = True
                else:
                    # Clean markdown links
                    import re
                    clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                    doc.add_paragraph(clean_line)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_notes_and_quiz(plan_content, subject, topic, grade):
        """Generate comprehensive notes and quiz from lesson plan"""
        prompt = f"""Based on the following lesson plan for {subject} - {topic} (Grade/Level: {grade}), generate:

1. **Comprehensive Study Notes** - Detailed notes that students can use for studying, including:
   - Key concepts and definitions
   - Important points and explanations
   - Examples and illustrations
   - Summary of main topics
   - Important formulas/theorems (if applicable)

2. **Quiz/Assessment Questions** - Create a quiz with:
   - 10-15 multiple choice questions (with 4 options each)
   - 5-7 short answer questions
   - 2-3 essay/long answer questions
   - Answer key with explanations

Format the response in Markdown with clear sections:
- # Study Notes
- # Quiz Questions
- # Answer Key

Make it appropriate for {grade} level students.

Lesson Plan Content:
{plan_content[:2000]}  # Limit to avoid token issues

Generate comprehensive, well-structured notes and quiz questions."""
        
        return LLM_Setup(prompt)

# --- Session State Initialization ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = None
if "current_plan" not in st.session_state:
    st.session_state.current_plan = None
if "plan_history" not in st.session_state:
    st.session_state.plan_history = []

# --- Dark / Light Mode Toggle ---
# Initialize dark mode in session state
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False

dark_mode = st.session_state.dark_mode

# --- Define colors based on mode ---
bg_color = "#0f172a" if dark_mode else "#f8fafc"
text_color = "#f1f5f9" if dark_mode else "#1e293b"
card_bg_color = "#1e293b" if dark_mode else "#ffffff"
input_bg_color = "#334155" if dark_mode else "#ffffff"
input_text_color = "#f1f5f9" if dark_mode else "#1e293b"
button_color = "#6366f1" if dark_mode else "#4f46e5"
accent_color = "#8b5cf6" if dark_mode else "#7c3aed"
border_color = "#475569" if dark_mode else "#e2e8f0"

# --- Enhanced CSS Styling ---
st.markdown(
    f"""
    <style>
    /* Main App Background */
    .stApp {{
        background: linear-gradient(135deg, {bg_color} 0%, {'#1e293b' if dark_mode else '#e0e7ff'} 100%);
        color: {text_color};
    }}
    
    /* Hide Streamlit default elements */
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
    header {{visibility: hidden;}}
    
    /* Custom Header */
    .main-header {{
        background: linear-gradient(135deg, {button_color} 0%, {accent_color} 100%);
        padding: 2rem 1rem;
        border-radius: 15px;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        text-align: center;
    }}
    
    .main-header h1 {{
        color: white;
        font-size: 2.5rem;
        font-weight: 800;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }}
    
    .main-header p {{
        color: rgba(255,255,255,0.9);
        font-size: 1.1rem;
        margin-top: 0.5rem;
    }}
    
    /* Cards */
    .lesson-card {{
        background-color: {card_bg_color};
        padding: 2rem;
        border-radius: 15px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border: 1px solid {border_color};
    }}
    
    /* Input boxes - Mobile Friendly */
    .stTextInput>div>div>input, 
    .stTextArea>div>div>textarea,
    .stSelectbox>div>div>select {{
        background-color: {input_bg_color} !important;
        color: {input_text_color} !important;
        font-size: 16px !important; /* Larger for mobile */
        border-radius: 10px !important;
        padding: 12px 15px !important;
        border: 2px solid {border_color} !important;
        transition: all 0.3s ease !important;
    }}
    
    .stTextInput>div>div>input:focus,
    .stTextArea>div>div>textarea:focus {{
        border-color: {button_color} !important;
        box-shadow: 0 0 0 3px rgba(99, 102, 241, 0.1) !important;
    }}
    
    /* Labels */
    .stTextInput label, 
    .stTextArea label,
    .stSelectbox label {{
        color: {text_color} !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        margin-bottom: 8px !important;
    }}
    
    /* Buttons - Touch Friendly */
    .stButton>button {{
        background: linear-gradient(135deg, {button_color} 0%, {accent_color} 100%) !important;
        color: white !important;
        font-weight: 600 !important;
        font-size: 16px !important;
        border-radius: 12px !important;
        padding: 12px 30px !important;
        min-height: 48px !important; /* Touch-friendly size */
        width: 100% !important;
        border: none !important;
        box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3) !important;
        transition: all 0.3s ease !important;
    }}
    
    .stButton>button:hover {{
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(99, 102, 241, 0.4) !important;
    }}
    
    .stButton>button:active {{
        transform: translateY(0) !important;
    }}
    
    /* Sidebar */
    .css-1d391kg {{
        background-color: {card_bg_color} !important;
    }}
    
    /* Markdown content */
    .lesson-plan-content {{
        background-color: {card_bg_color};
        padding: 2rem;
        border-radius: 15px;
        margin: 1rem 0;
        line-height: 1.8;
        font-size: 16px;
    }}
    
    .lesson-plan-content h1, .lesson-plan-content h2, .lesson-plan-content h3 {{
        color: {button_color};
        margin-top: 1.5rem;
        margin-bottom: 1rem;
    }}
    
    .lesson-plan-content a {{
        color: {accent_color};
        text-decoration: none;
        font-weight: 600;
    }}
    
    .lesson-plan-content a:hover {{
        text-decoration: underline;
    }}
    
    /* Mobile Responsive */
    @media (max-width: 768px) {{
        .main-header h1 {{
            font-size: 1.8rem;
        }}
        
        .lesson-card {{
            padding: 1.5rem;
        }}
        
        [data-testid="stSidebar"] {{
            width: 100% !important;
        }}
    }}
    
    /* Success/Error Messages */
    .stSuccess {{
        background-color: rgba(34, 197, 94, 0.1);
        border-left: 4px solid #22c55e;
        padding: 1rem;
        border-radius: 8px;
    }}
    
    .stError {{
        background-color: rgba(239, 68, 68, 0.1);
        border-left: 4px solid #ef4444;
        padding: 1rem;
        border-radius: 8px;
    }}
    
    /* Loading Spinner */
    .stSpinner > div {{
        border-top-color: {button_color} !important;
    }}
    
    /* Navigation Bar */
    .navbar {{
        background: linear-gradient(135deg, {button_color} 0%, {accent_color} 100%);
        padding: 1rem 2rem;
        border-radius: 0 0 15px 15px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        position: relative;
    }}
    
    .navbar-content {{
        display: flex;
        justify-content: space-between;
        align-items: center;
        max-width: 100%;
        margin: 0 auto;
    }}
    
    .navbar-brand {{
        font-size: 1.5rem;
        font-weight: 800;
        color: white;
        text-decoration: none;
    }}
    
    .navbar-links {{
        display: flex;
        gap: 1.5rem;
        align-items: center;
    }}
    
    .navbar-link {{
        color: white;
        text-decoration: none;
        font-weight: 600;
        padding: 0.5rem 1rem;
        border-radius: 8px;
        transition: all 0.3s ease;
    }}
    
    .navbar-link:hover {{
        background: rgba(255,255,255,0.2);
    }}
    
    /* Dark Mode Toggle Button in Navbar */
    .dark-mode-toggle {{
        background: rgba(255, 255, 255, 0.15) !important;
        color: white !important;
        border: 2px solid rgba(255, 255, 255, 0.3) !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        padding: 8px 16px !important;
        border-radius: 25px !important;
        transition: all 0.3s ease !important;
        cursor: pointer !important;
        display: flex !important;
        align-items: center !important;
        gap: 8px !important;
    }}
    
    .dark-mode-toggle:hover {{
        background: rgba(255, 255, 255, 0.25) !important;
        transform: scale(1.05) !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.2) !important;
    }}
    
    .dark-mode-toggle:active {{
        transform: scale(0.98) !important;
    }}
    
    /* Toggle Icon Animation */
    .toggle-icon {{
        font-size: 18px !important;
        transition: transform 0.3s ease !important;
    }}
    
    .dark-mode-toggle:hover .toggle-icon {{
        transform: rotate(15deg) !important;
    }}
    
    /* Footer */
    .footer {{
        background: linear-gradient(135deg, {card_bg_color} 0%, {bg_color} 100%);
        padding: 2rem;
        margin-top: 4rem;
        border-top: 2px solid {border_color};
        border-radius: 15px 15px 0 0;
    }}
    
    .footer-content {{
        max-width: 1200px;
        margin: 0 auto;
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
        gap: 2rem;
    }}
    
    .footer-section h4 {{
        color: {button_color};
        margin-bottom: 1rem;
        font-size: 1.2rem;
    }}
    
    .footer-section p, .footer-section a {{
        color: {text_color};
        text-decoration: none;
        line-height: 1.8;
    }}
    
    .footer-section a:hover {{
        color: {accent_color};
    }}
    
    .footer-bottom {{
        text-align: center;
        margin-top: 2rem;
        padding-top: 1rem;
        border-top: 1px solid {border_color};
        color: {text_color};
    }}
    
    /* Feature Cards */
    .feature-card {{
        background: {card_bg_color};
        padding: 2rem;
        border-radius: 15px;
        border: 2px solid {border_color};
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        transition: transform 0.3s ease, box-shadow 0.3s ease;
        height: 100%;
    }}
    
    .feature-card:hover {{
        transform: translateY(-5px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.15);
    }}
    
    .feature-icon {{
        font-size: 3rem;
        margin-bottom: 1rem;
    }}
    
    .feature-card h3 {{
        color: {button_color};
        margin-bottom: 1rem;
    }}
    
    .feature-card p {{
        color: {text_color};
        line-height: 1.6;
    }}
    </style>
    """,
    unsafe_allow_html=True
)

# --- Page Header ---
st.markdown(
    f"""
    <div class="main-header">
        <h1>📚 AI Lesson Planner</h1>
        <p>Create engaging, interactive lesson plans with AI-powered assistance</p>
    </div>
    """,
    unsafe_allow_html=True
)

# --- Login / Signup Section ---
if not st.session_state.logged_in:
    # Dark Mode Toggle for Login Page
    with st.sidebar:
        st.markdown("### 🎨 Theme")
        dark_mode_toggle_login = st.toggle(
            "🌙 Dark Mode",
            value=dark_mode,
            key="dark_mode_toggle_login"
        )
        if dark_mode_toggle_login != dark_mode:
            st.session_state.dark_mode = dark_mode_toggle_login
            st.rerun()
        st.markdown("---")
    
    # Main content area with tabs for better visibility
    tab1, tab2 = st.tabs(["🔐 Login", "✨ Sign Up"])
    
    with tab1:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown(
                f"<div class='lesson-card'><h3 style='color: {text_color}; text-align: center; margin-bottom: 1.5rem;'>Welcome Back</h3>",
                unsafe_allow_html=True
            )
            username = st.text_input("👤 Username", key="login_username", placeholder="Enter your username")
            password = st.text_input("🔒 Password", type="password", key="login_password", placeholder="Enter your password")
            if st.button("🔑 Login", key="login_btn", use_container_width=True):
                user = users.find_one({"username": username})
                if user and bcrypt.checkpw(password.encode('utf-8'), user["password"]):
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.success(f"✅ Welcome back, {username}!")
                    st.rerun()
                else:
                    st.error("❌ Invalid username or password")
            st.markdown("</div>", unsafe_allow_html=True)
    
    with tab2:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown(
                f"<div class='lesson-card'><h3 style='color: {text_color}; text-align: center; margin-bottom: 1.5rem;'>Create Your Account</h3>",
                unsafe_allow_html=True
            )
            username = st.text_input("👤 Username", key="signup_username", placeholder="Enter your username")
            password = st.text_input("🔒 Password", type="password", key="signup_password", placeholder="Enter your password (min. 6 characters)")
            confirm_password = st.text_input("🔒 Confirm Password", type="password", key="confirm_password", placeholder="Re-enter your password")
            if st.button("🚀 Create Account", key="create_account_btn", use_container_width=True):
                if not username or not password or not confirm_password:
                    st.warning("⚠️ Please fill all fields")
                elif len(password) < 6:
                    st.warning("⚠️ Password must be at least 6 characters long")
                elif password != confirm_password:
                    st.error("⚠️ Passwords do not match. Please try again.")
                elif users.find_one({"username": username}):
                    st.error("⚠️ Username already exists. Try another.")
                else:
                    hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
                    users.insert_one({"username": username, "password": hashed})
                    st.success("✅ Account created successfully! Please switch to the Login tab to sign in.")
            st.markdown("</div>", unsafe_allow_html=True)

# --- AI Lesson Planner Section (Only if logged in) ---
if st.session_state.logged_in:
    # Navigation Bar with Dark Mode Toggle
    # Create a container for navbar with toggle
    nav_col1, nav_col2 = st.columns([5, 1])
    
    with nav_col1:
        st.markdown(
            f"""
            <div class="navbar">
                <div class="navbar-content">
                    <div class="navbar-brand">📚 AI Lesson Planner</div>
                    <div class="navbar-links">
                        <a href="#" class="navbar-link">Home</a>
                        <a href="#" class="navbar-link">Features</a>
                        <a href="#" class="navbar-link">About</a>
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    
    with nav_col2:
        # Dark Mode Toggle Button - Enhanced UI
        toggle_icon = "🌙" if not dark_mode else "☀️"
        toggle_text = "Dark" if not dark_mode else "Light"
        
        # Add some spacing
        st.markdown("<br>", unsafe_allow_html=True)
        
        if st.button(
            f"{toggle_icon} {toggle_text}",
            key="dark_mode_toggle_navbar",
            use_container_width=True,
            help="Toggle between dark and light mode"
        ):
            st.session_state.dark_mode = not st.session_state.dark_mode
            st.rerun()
        
        # Enhanced custom styling to the toggle button
        st.markdown(
            f"""
            <style>
            button[key="dark_mode_toggle_navbar"] {{
                background: linear-gradient(135deg, rgba(255, 255, 255, 0.2) 0%, rgba(255, 255, 255, 0.1) 100%) !important;
                color: white !important;
                border: 2px solid rgba(255, 255, 255, 0.4) !important;
                font-weight: 700 !important;
                font-size: 15px !important;
                border-radius: 30px !important;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
                margin-top: 0.5rem !important;
                padding: 10px 20px !important;
                box-shadow: 0 2px 8px rgba(0,0,0,0.15) !important;
                text-transform: none !important;
                letter-spacing: 0.5px !important;
            }}
            
            button[key="dark_mode_toggle_navbar"]:hover {{
                background: linear-gradient(135deg, rgba(255, 255, 255, 0.3) 0%, rgba(255, 255, 255, 0.2) 100%) !important;
                transform: translateY(-2px) scale(1.05) !important;
                box-shadow: 0 6px 20px rgba(0,0,0,0.25) !important;
                border-color: rgba(255, 255, 255, 0.6) !important;
            }}
            
            button[key="dark_mode_toggle_navbar"]:active {{
                transform: translateY(0) scale(0.98) !important;
                box-shadow: 0 2px 8px rgba(0,0,0,0.15) !important;
            }}
            </style>
            """,
            unsafe_allow_html=True
        )
    
    # Sidebar
    with st.sidebar:
        st.markdown(f"### 👤 {st.session_state.username}")
        st.markdown("---")
        
        # Dark Mode Toggle in Sidebar (alternative option)
        dark_mode_toggle_sidebar = st.toggle(
            "🌙 Dark Mode" if not dark_mode else "☀️ Light Mode",
            value=dark_mode,
            key="dark_mode_toggle_sidebar"
        )
        if dark_mode_toggle_sidebar != dark_mode:
            st.session_state.dark_mode = dark_mode_toggle_sidebar
            st.rerun()
        
        st.markdown("---")
        
        # Navigation
        nav_options = ["🏠 Home", "📝 Create Plan", "📚 My Plans", "⚙️ Settings"]
        if "nav_page" in st.session_state and st.session_state.nav_page in nav_options:
            default_index = nav_options.index(st.session_state.nav_page)
        else:
            default_index = 0
        page = st.radio("📑 Navigation", nav_options, index=default_index, label_visibility="collapsed")
        if "nav_page" in st.session_state:
            del st.session_state.nav_page
        
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.username = None
            st.session_state.current_plan = None
            st.rerun()
    
    # Home Page
    if page == "🏠 Home":
        st.markdown("### 👋 Welcome Back, {}!".format(st.session_state.username))
        st.markdown("---")
        
        # Welcome Section
        col_welcome1, col_welcome2 = st.columns([2, 1])
        with col_welcome1:
             st.markdown(
                f"""
                <div style='background: linear-gradient(135deg, {button_color} 0%, {accent_color} 100%); padding: 2rem; border-radius: 15px; color: white; margin-bottom: 2rem;'>
                    <h2 style='color: white; margin-bottom: 1rem;'>🚀 Get Started with AI Lesson Planning</h2>
                    <p style='color: rgba(255,255,255,0.9); font-size: 1.1rem; line-height: 1.8;'>
                        Create comprehensive, engaging lesson plans in minutes with AI-powered assistance. 
                        Generate detailed plans with YouTube links, study notes, quizzes, and more!
                    </p>
                </div>
                """,
        unsafe_allow_html=True
    )

        with col_welcome2:
            try:
                saved_plans = list(lesson_plans.find({"username": st.session_state.username}))
                recent_plans = sorted(saved_plans, key=lambda x: x.get('created_at', datetime.min) if isinstance(x.get('created_at'), datetime) else datetime.min, reverse=True)[:3]
                
                st.markdown("### 📊 Quick Stats")
                st.metric("Total Plans", len(saved_plans))
                st.metric("This Week", len([p for p in saved_plans if isinstance(p.get('created_at'), datetime) and (datetime.now() - p['created_at']).days <= 7]))
                
                if recent_plans:
                    st.markdown("### 📚 Recent Plans")
                    for plan in recent_plans[:3]:
                        st.markdown(f"• **{plan['subject']}** - {plan['topic']}")
            except:
                pass
        
        # Quick Actions
        st.markdown("### ⚡ Quick Actions")
        col_action1, col_action2, col_action3 = st.columns(3)
        
        with col_action1:
            if st.button("📝 Create New Plan", use_container_width=True, type="primary"):
                st.session_state.nav_page = "📝 Create Plan"
                st.rerun()
        
        with col_action2:
            if st.button("📚 View My Plans", use_container_width=True):
                st.session_state.nav_page = "📚 My Plans"
                st.rerun()
        
        with col_action3:
            if st.button("📊 View Statistics", use_container_width=True):
                st.session_state.nav_page = "📚 My Plans"
                st.rerun()
        
        st.markdown("---")
        
        # Features Section
        st.markdown("### ✨ Key Features")
        col_feat1, col_feat2, col_feat3 = st.columns(3)
        
        with col_feat1:
            st.markdown(
                f"""
                <div class="feature-card">
                    <div class="feature-icon">🤖</div>
                    <h3>AI-Powered Generation</h3>
                    <p>Generate comprehensive lesson plans with AI assistance. Includes YouTube links, activities, and assessments tailored to your needs.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        with col_feat2:
            st.markdown(
                f"""
                <div class="feature-card">
                    <div class="feature-icon">📚</div>
                    <h3>Notes & Quiz Generator</h3>
                    <p>Automatically generate study notes and quizzes from your lesson plans. Perfect for student materials and assessments.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        with col_feat3:
            st.markdown(
                f"""
                <div class="feature-card">
                    <div class="feature-icon">📥</div>
                    <h3>Multiple Export Formats</h3>
                    <p>Export your plans in PDF, Word, or Markdown format. Share with colleagues or print for classroom use.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        col_feat4, col_feat5, col_feat6 = st.columns(3)
        
        with col_feat4:
            st.markdown(
                f"""
                <div class="feature-card">
                    <div class="feature-icon">🎓</div>
                    <h3>All Education Levels</h3>
                    <p>From Kindergarten to PhD, create lesson plans for any educational level. Customize difficulty and learning styles.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        with col_feat5:
            st.markdown(
                f"""
                <div class="feature-card">
                    <div class="feature-icon">🔍</div>
                    <h3>Smart Search & Filter</h3>
                    <p>Easily find your plans with advanced search and filter options. Sort by date, subject, or grade level.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        with col_feat6:
            st.markdown(
                f"""
                <div class="feature-card">
                    <div class="feature-icon">💾</div>
                    <h3>Cloud Storage</h3>
                    <p>All your plans are safely stored in the cloud. Access them anytime, anywhere, from any device.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        st.markdown("---")
        
        # How It Works
        st.markdown("### 🎯 How It Works")
        st.markdown(
            """
            <div style='background-color: {}; padding: 2rem; border-radius: 15px; margin: 1rem 0;'>
                <div style='display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 2rem;'>
                    <div style='text-align: center;'>
                        <div style='font-size: 3rem; margin-bottom: 1rem;'>1️⃣</div>
                        <h4 style='color: {};'>Fill Details</h4>
                        <p style='color: {};'>Enter subject, topic, grade level, and learning objectives</p>
                    </div>
                    <div style='text-align: center;'>
                        <div style='font-size: 3rem; margin-bottom: 1rem;'>2️⃣</div>
                        <h4 style='color: {};'>AI Generation</h4>
                        <p style='color: {};'>AI creates a comprehensive lesson plan with resources</p>
                    </div>
                    <div style='text-align: center;'>
                        <div style='font-size: 3rem; margin-bottom: 1rem;'>3️⃣</div>
                        <h4 style='color: {};'>Customize & Export</h4>
                        <p style='color: {};'>Review, customize, and export in your preferred format</p>
                    </div>
                </div>
            </div>
            """.format(card_bg_color, button_color, text_color, button_color, text_color, button_color, text_color),
            unsafe_allow_html=True
        )
        
        # Footer will be added at the end
    
    # Create Plan Page
    elif page == "📝 Create Plan":
        st.markdown("### 🎯 Create New Lesson Plan")
        st.markdown("Fill in the details below to generate a comprehensive lesson plan with YouTube links and resources.")
        
        # Form in columns for better layout
        col1, col2 = st.columns(2)
        
        with col1:
            subject = st.text_input('📖 Subject *', key="subject", placeholder="e.g., Science, Mathematics, English")
            topic = st.text_input('📌 Topic *', key="topic", placeholder="e.g., Photosynthesis, Algebra, Poetry")
            grade = st.selectbox('🎓 Grade/Level *', 
                                ['Kindergarten', 'Grade 1', 'Grade 2', 'Grade 3', 'Grade 4', 'Grade 5',
                                 'Grade 6', 'Grade 7', 'Grade 8', 'Grade 9', 'Grade 10', 'Grade 11', 'Grade 12',
                                 'Associate Degree', 'Bachelor\'s Degree', 'Master\'s Degree', 'PhD/Doctorate', 'Professional Development'],
                                key="grade")
        
        with col2:
            duration = st.text_input('⏱️ Duration *', key="duration", placeholder="e.g., 45 minutes, 1 hour")
            learning_style = st.selectbox('🎨 Learning Style', 
                                         ['Visual', 'Auditory', 'Kinesthetic', 'Mixed', 'Interactive'],
                                         key="learning_style")
            difficulty = st.selectbox('📊 Difficulty Level', 
                                     ['Beginner', 'Intermediate', 'Advanced'],
                                     key="difficulty")
        
        learning_objectives = st.text_area('🎯 Learning Objectives *', 
                                          key="learning_objectives",
                                          placeholder="What should students learn? (e.g., Understand the process of photosynthesis, Identify key components)")
        customization = st.text_area('✨ Additional Customization (Optional)', 
                                    key="customization",
                                    placeholder="Any specific requirements? (e.g., Include group activities, Focus on hands-on experiments)")
        
        # Generate button
        col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
        with col_btn2:
            if st.button('🚀 Generate Lesson Plan', use_container_width=True, type="primary"):
                if not subject or not topic or not grade or not duration or not learning_objectives:
                    st.warning('⚠️ Please fill out all required fields (marked with *) before generating the lesson plan.')
                else:
                    with st.spinner("🧠 AI is crafting your comprehensive lesson plan with YouTube links and resources..."):
                        # Enhanced prompt with YouTube links requirement
                        prompt = f"""Generate a comprehensive, detailed lesson plan for the subject "{subject}" on the topic "{topic}".

**Lesson Details:**
- Grade Level: {grade}
- Duration: {duration}
- Learning Style: {learning_style}
- Difficulty Level: {difficulty}

**Learning Objectives:**
{learning_objectives}

**Customization Requirements:**
{customization if customization else "None specified"}

**Requirements:**
1. Create a well-structured lesson plan in Markdown format
2. Include relevant YouTube video links (at least 2-3 videos) that are educational and appropriate for {grade} level
3. Format YouTube links as: [Video Title](https://www.youtube.com/watch?v=VIDEO_ID) or [Video Title](https://youtu.be/VIDEO_ID)
4. Include the following sections:
   - **Lesson Overview** (brief summary)
   - **Learning Objectives** (detailed, measurable, using Bloom's taxonomy for higher education levels)
   - **Materials Needed** (list all resources, including digital tools if applicable)
   - **Introduction/Warm-up** (5-10 minutes, or appropriate for session length)
   - **Main Content** (detailed step-by-step activities with timing)
   - **YouTube Videos & Resources** (with descriptions of what each video covers and why it's relevant)
   - **Interactive Activities** (hands-on, group activities, discussions, or case studies)
   - **Assessment/Evaluation** (how to measure learning - quizzes, assignments, projects, presentations)
   - **Homework/Extension Activities** (optional follow-up work or research)
   - **Additional Resources** (websites, articles, research papers, academic journals if applicable)

5. Make it engaging, interactive, and appropriate for {grade} level:
   - For K-12: Use age-appropriate language, include games and hands-on activities
   - For Associate/Bachelor's: Include academic rigor, research components, and critical thinking
   - For Master's/PhD: Focus on advanced concepts, research methodologies, scholarly discussions, and peer review
   - For Professional Development: Emphasize practical applications, real-world scenarios, and skill-building
6. Include specific time allocations for each section
7. Add practical examples and real-world connections relevant to the level
8. Ensure the content aligns with {learning_style} learning style
9. Adjust complexity, depth, and academic rigor based on {difficulty} difficulty level and {grade} level
10. For higher education levels, include:
    - Academic citations and references where appropriate
    - Discussion questions that promote critical thinking
    - Research assignments or literature reviews
    - Peer collaboration and presentation opportunities

Return the lesson plan in clean Markdown format with proper headings, bullet points, and formatting."""
                        
                        try:
                            llm_output = LLM_Setup(prompt)
                            st.session_state.current_plan = {
                                "subject": subject,
                                "topic": topic,
                                "grade": grade,
                                "duration": duration,
                                "content": llm_output,
                                "created_at": datetime.now().isoformat()
                            }
                            st.success("✅ Lesson Plan Generated Successfully!")
                        except Exception as e:
                            st.error(f"❌ Error generating lesson plan: {str(e)}")
                            st.info("💡 Please try again. If the issue persists, check your API key and internet connection.")
        
        # Display generated plan
        if st.session_state.current_plan:
            st.markdown("---")
            st.markdown("### 📄 Your Lesson Plan")
            
            # Action buttons - Row 1
            col_save, col_notes, col_clear = st.columns(3)
            with col_save:
                if st.button("💾 Save Plan", use_container_width=True):
                    try:
                        plan_data = {
                            "username": st.session_state.username,
                            "subject": st.session_state.current_plan["subject"],
                            "topic": st.session_state.current_plan["topic"],
                            "grade": st.session_state.current_plan["grade"],
                            "duration": st.session_state.current_plan["duration"],
                            "content": st.session_state.current_plan["content"],
                            "created_at": datetime.now()
                        }
                        lesson_plans.insert_one(plan_data)
                        st.success("✅ Lesson plan saved successfully!")
                    except Exception as e:
                        st.error(f"❌ Error saving plan: {str(e)}")
            
            with col_notes:
                if st.button("📝 Generate Notes & Quiz", use_container_width=True):
                    with st.spinner("🧠 Generating comprehensive notes and quiz..."):
                        try:
                            notes_quiz = generate_notes_and_quiz(
                                st.session_state.current_plan['content'],
                                st.session_state.current_plan['subject'],
                                st.session_state.current_plan['topic'],
                                st.session_state.current_plan['grade']
                            )
                            st.session_state.notes_quiz = notes_quiz
                            st.success("✅ Notes and Quiz generated!")
                        except Exception as e:
                            st.error(f"❌ Error generating notes/quiz: {str(e)}")
            
            with col_clear:
                if st.button("🗑️ Clear", use_container_width=True):
                    st.session_state.current_plan = None
                    if "notes_quiz" in st.session_state:
                        del st.session_state.notes_quiz
                    st.rerun()
            
            # Export buttons - Row 2
            st.markdown("### 📤 Export Options")
            col_pdf, col_word, col_md = st.columns(3)
            
            plan_data_export = {
                "subject": st.session_state.current_plan['subject'],
                "topic": st.session_state.current_plan['topic'],
                "grade": st.session_state.current_plan['grade'],
                "duration": st.session_state.current_plan['duration'],
                "content": st.session_state.current_plan['content'],
                "created_at": datetime.now().strftime('%Y-%m-%d %H:%M')
            }
            
            with col_pdf:
                if REPORTLAB_AVAILABLE:
                    try:
                        pdf_data = generate_pdf(plan_data_export)
                        if pdf_data:
                            st.download_button(
                                label="📄 Download PDF",
                                data=pdf_data,
                                file_name=f"lesson_plan_{st.session_state.current_plan['subject']}_{st.session_state.current_plan['topic']}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                        else:
                            st.button("📄 PDF (Unavailable)", disabled=True, use_container_width=True)
                    except Exception as e:
                        st.error(f"PDF Error: {str(e)}")
                else:
                    st.button("📄 PDF (Install reportlab)", disabled=True, use_container_width=True)
            
            with col_word:
                if DOCX_AVAILABLE:
                    try:
                        word_data = generate_word_doc(plan_data_export)
                        if word_data:
                            st.download_button(
                                label="📘 Download Word",
                                data=word_data,
                                file_name=f"lesson_plan_{st.session_state.current_plan['subject']}_{st.session_state.current_plan['topic']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            st.button("📘 Word (Unavailable)", disabled=True, use_container_width=True)
                    except Exception as e:
                        st.error(f"Word Error: {str(e)}")
                else:
                    st.button("📘 Word (Install python-docx)", disabled=True, use_container_width=True)
            
            with col_md:
                plan_markdown = f"""# {st.session_state.current_plan['subject']} - {st.session_state.current_plan['topic']}

**Grade/Level:** {st.session_state.current_plan['grade']}  
**Duration:** {st.session_state.current_plan['duration']}  
**Created:** {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

{st.session_state.current_plan['content']}
"""
                st.download_button(
                    label="📝 Download Markdown",
                    data=plan_markdown,
                    file_name=f"lesson_plan_{st.session_state.current_plan['subject']}_{st.session_state.current_plan['topic']}.md",
                    mime="text/markdown",
                    use_container_width=True
                )
            
            # Display Notes & Quiz if generated
            if "notes_quiz" in st.session_state and st.session_state.notes_quiz:
                st.markdown("---")
                st.markdown("### 📚 Generated Study Notes & Quiz")
                st.markdown(
                    f"<div class='lesson-plan-content'>{st.session_state.notes_quiz}</div>",
                    unsafe_allow_html=True
                )
                
                # Download notes & quiz
                col_notes_pdf, col_notes_word, col_notes_md = st.columns(3)
                notes_data = {
                    "subject": f"{st.session_state.current_plan['subject']} - Notes & Quiz",
                    "topic": st.session_state.current_plan['topic'],
                    "grade": st.session_state.current_plan['grade'],
                    "duration": "N/A",
                    "content": st.session_state.notes_quiz,
                    "created_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                }
                
                with col_notes_pdf:
                    if REPORTLAB_AVAILABLE:
                        try:
                            pdf_data = generate_pdf(notes_data)
                            if pdf_data:
                                st.download_button(
                                    label="📄 Download Notes PDF",
                                    data=pdf_data,
                                    file_name=f"notes_quiz_{st.session_state.current_plan['subject']}_{st.session_state.current_plan['topic']}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                        except:
                            pass
                
                with col_notes_word:
                    if DOCX_AVAILABLE:
                        try:
                            word_data = generate_word_doc(notes_data)
                            if word_data:
                                st.download_button(
                                    label="📘 Download Notes Word",
                                    data=word_data,
                                    file_name=f"notes_quiz_{st.session_state.current_plan['subject']}_{st.session_state.current_plan['topic']}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                        except:
                            pass
                
                with col_notes_md:
                    st.download_button(
                        label="📝 Download Notes MD",
                        data=st.session_state.notes_quiz,
                        file_name=f"notes_quiz_{st.session_state.current_plan['subject']}_{st.session_state.current_plan['topic']}.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
            
            # Display plan content
            st.markdown(
                f"<div class='lesson-plan-content'>{st.session_state.current_plan['content']}</div>",
                unsafe_allow_html=True
            )

    # My Plans Page
    elif page == "📚 My Plans":
        st.markdown("### 📚 My Saved Lesson Plans")
        st.markdown("Manage and organize all your lesson plans in one place")
        
        try:
            saved_plans = list(lesson_plans.find(
                {"username": st.session_state.username}
            ).sort("created_at", -1))
            
            if saved_plans:
                # Enhanced Statistics Cards
                st.markdown("---")
                col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
                with col_stat1:
                    st.markdown(
                        f"""
                        <div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 1.5rem; border-radius: 15px; text-align: center; color: white;'>
                            <h2 style='margin: 0; font-size: 2.5rem;'>{len(saved_plans)}</h2>
                            <p style='margin: 0.5rem 0 0 0; font-size: 1rem;'>📊 Total Plans</p>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                with col_stat2:
                    subjects = set([p['subject'] for p in saved_plans])
                    st.markdown(
                        f"""
                        <div style='background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 1.5rem; border-radius: 15px; text-align: center; color: white;'>
                            <h2 style='margin: 0; font-size: 2.5rem;'>{len(subjects)}</h2>
                            <p style='margin: 0.5rem 0 0 0; font-size: 1rem;'>📖 Subjects</p>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                with col_stat3:
                    grades = set([p['grade'] for p in saved_plans])
                    st.markdown(
                        f"""
                        <div style='background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 1.5rem; border-radius: 15px; text-align: center; color: white;'>
                            <h2 style='margin: 0; font-size: 2.5rem;'>{len(grades)}</h2>
                            <p style='margin: 0.5rem 0 0 0; font-size: 1rem;'>🎓 Levels</p>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                with col_stat4:
                    recent_plans = [p for p in saved_plans if isinstance(p.get('created_at'), datetime) and (datetime.now() - p['created_at']).days <= 7]
                    st.markdown(
                        f"""
                        <div style='background: linear-gradient(135deg, #43e97b 0%, #38f9d7 100%); padding: 1.5rem; border-radius: 15px; text-align: center; color: white;'>
                            <h2 style='margin: 0; font-size: 2.5rem;'>{len(recent_plans)}</h2>
                            <p style='margin: 0.5rem 0 0 0; font-size: 1rem;'>🆕 This Week</p>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                
                st.markdown("---")
                
                # Enhanced Search and Filter
                st.markdown("### 🔍 Search & Filter")
                col_search, col_filter, col_sort = st.columns([2, 1, 1])
                with col_search:
                    search_query = st.text_input("🔍 Search plans", placeholder="Search by subject, topic, or grade...", key="search_plans", label_visibility="collapsed")
                with col_filter:
                    filter_grade = st.selectbox("🎓 Filter by Level", 
                                               ["All Levels"] + sorted(set([p['grade'] for p in saved_plans])),
                                               key="filter_grade", label_visibility="collapsed")
                with col_sort:
                    sort_option = st.selectbox("📊 Sort by", 
                                              ["Newest First", "Oldest First", "Subject A-Z", "Subject Z-A"],
                                              key="sort_plans", label_visibility="collapsed")
                
                # Filter plans
                filtered_plans = saved_plans
                if search_query:
                    search_lower = search_query.lower()
                    filtered_plans = [p for p in filtered_plans 
                                     if search_lower in p['subject'].lower() 
                                     or search_lower in p['topic'].lower() 
                                     or search_lower in p['grade'].lower()]
                
                if filter_grade != "All Levels":
                    filtered_plans = [p for p in filtered_plans if p['grade'] == filter_grade]
                
                # Sort plans
                if sort_option == "Newest First":
                    filtered_plans = sorted(filtered_plans, key=lambda x: x.get('created_at', datetime.min) if isinstance(x.get('created_at'), datetime) else datetime.min, reverse=True)
                elif sort_option == "Oldest First":
                    filtered_plans = sorted(filtered_plans, key=lambda x: x.get('created_at', datetime.max) if isinstance(x.get('created_at'), datetime) else datetime.max)
                elif sort_option == "Subject A-Z":
                    filtered_plans = sorted(filtered_plans, key=lambda x: x.get('subject', '').lower())
                elif sort_option == "Subject Z-A":
                    filtered_plans = sorted(filtered_plans, key=lambda x: x.get('subject', '').lower(), reverse=True)
                
                if filtered_plans:
                    st.markdown(f"**📋 Showing {len(filtered_plans)} of {len(saved_plans)} plans**")
                    st.markdown("---")
                    
                    # Modern Card Layout
                    for idx, plan in enumerate(filtered_plans):
                        plan_date = plan['created_at'].strftime('%B %d, %Y at %I:%M %p') if isinstance(plan['created_at'], datetime) else str(plan.get('created_at', 'Unknown'))
                        days_ago = (datetime.now() - plan['created_at']).days if isinstance(plan['created_at'], datetime) else 0
                        
                        # Card styling
                        st.markdown(
                            f"""
                            <div style='background-color: {card_bg_color}; padding: 1.5rem; border-radius: 15px; margin: 1rem 0; border: 2px solid {border_color}; box-shadow: 0 4px 15px rgba(0,0,0,0.1);'>
                                <div style='display: flex; justify-content: space-between; align-items: start; margin-bottom: 1rem;'>
                                    <div>
                                        <h3 style='color: {button_color}; margin: 0; font-size: 1.5rem;'>{plan['subject']}</h3>
                                        <p style='color: {text_color}; margin: 0.5rem 0; font-size: 1.1rem; font-weight: 600;'>{plan['topic']}</p>
                                    </div>
                                    <div style='text-align: right;'>
                                        <span style='background: {accent_color}; color: white; padding: 0.3rem 0.8rem; border-radius: 20px; font-size: 0.85rem;'>{plan['grade']}</span>
                                    </div>
                                </div>
                                <div style='display: flex; gap: 2rem; margin-bottom: 1rem; color: {text_color};'>
                                    <div>⏱️ <strong>Duration:</strong> {plan['duration']}</div>
                                    <div>📅 <strong>Created:</strong> {plan_date}</div>
                                    {f"<div>🕐 <strong>{days_ago} days ago</strong></div>" if days_ago > 0 else ""}
                                </div>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )
                        
                        # Action buttons in columns
                        col_view, col_duplicate, col_export, col_delete = st.columns(4)
                        
                        with col_view:
                            if st.button("👁️ View Full Plan", key=f"view_{idx}", use_container_width=True):
                                st.session_state.current_plan = {
                                    "subject": plan['subject'],
                                    "topic": plan['topic'],
                                    "grade": plan['grade'],
                                    "duration": plan['duration'],
                                    "content": plan['content'],
                                    "created_at": plan['created_at'].isoformat() if isinstance(plan['created_at'], datetime) else str(plan.get('created_at', ''))
                                }
                                st.rerun()
                        
                        with col_duplicate:
                            if st.button("📋 Duplicate", key=f"duplicate_{idx}", use_container_width=True):
                                new_plan = {
                                    "username": st.session_state.username,
                                    "subject": plan['subject'] + " (Copy)",
                                    "topic": plan['topic'],
                                    "grade": plan['grade'],
                                    "duration": plan['duration'],
                                    "content": plan['content'],
                                    "created_at": datetime.now()
                                }
                                lesson_plans.insert_one(new_plan)
                                st.success("✅ Plan duplicated!")
                                st.rerun()
                        
                        with col_export:
                            # Export dropdown
                            export_format = st.selectbox(
                                "📥 Export",
                                ["Markdown", "PDF", "Word"],
                                key=f"export_format_{idx}",
                                label_visibility="collapsed"
                            )
                            
                            plan_data_export = {
                                "subject": plan['subject'],
                                "topic": plan['topic'],
                                "grade": plan['grade'],
                                "duration": plan['duration'],
                                "content": plan['content'],
                                "created_at": plan_date
                            }
                            
                            if export_format == "Markdown":
                                plan_markdown = f"""# {plan['subject']} - {plan['topic']}

**Grade/Level:** {plan['grade']}  
**Duration:** {plan['duration']}  
**Created:** {plan_date}

---

{plan['content']}
"""
                                st.download_button(
                                    label="📝 Download",
                                    data=plan_markdown,
                                    file_name=f"lesson_plan_{plan['subject']}_{plan['topic']}.md",
                                    mime="text/markdown",
                                    key=f"download_md_{idx}",
                                    use_container_width=True
                                )
                            elif export_format == "PDF" and REPORTLAB_AVAILABLE:
                                try:
                                    pdf_data = generate_pdf(plan_data_export)
                                    if pdf_data:
                                        st.download_button(
                                            label="📄 Download",
                                            data=pdf_data,
                                            file_name=f"lesson_plan_{plan['subject']}_{plan['topic']}.pdf",
                                            mime="application/pdf",
                                            key=f"download_pdf_{idx}",
                                            use_container_width=True
                                        )
                                except:
                                    st.button("📄 PDF Error", disabled=True, use_container_width=True, key=f"pdf_err_{idx}")
                            elif export_format == "Word" and DOCX_AVAILABLE:
                                try:
                                    word_data = generate_word_doc(plan_data_export)
                                    if word_data:
                                        st.download_button(
                                            label="📘 Download",
                                            data=word_data,
                                            file_name=f"lesson_plan_{plan['subject']}_{plan['topic']}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"download_word_{idx}",
                                            use_container_width=True
                                        )
                                except:
                                    st.button("📘 Word Error", disabled=True, use_container_width=True, key=f"word_err_{idx}")
                        
                        with col_delete:
                            if st.button("🗑️ Delete", key=f"delete_{idx}", use_container_width=True):
                                lesson_plans.delete_one({"_id": plan["_id"]})
                                st.success("✅ Plan deleted!")
                                st.rerun()
                        
                        st.markdown("---")
                else:
                    st.info("🔍 No plans match your search criteria. Try adjusting your filters.")
            else:
                st.markdown(
                    """
                    <div style='text-align: center; padding: 3rem; background-color: #f8f9fa; border-radius: 15px; margin: 2rem 0;'>
                        <h2 style='color: #6c757d;'>📭 No Plans Yet</h2>
                        <p style='color: #6c757d; font-size: 1.1rem;'>You haven't saved any lesson plans yet.</p>
                        <p style='color: #6c757d;'>Create your first lesson plan to get started! 🚀</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
        except Exception as e:
            st.error(f"❌ Error loading plans: {str(e)}")
    
    # Settings Page
    elif page == "⚙️ Settings":
        st.markdown("### ⚙️ Settings")
        st.info("💡 Settings and preferences will be available here in future updates.")
        
        col_set1, col_set2 = st.columns(2)
        with col_set1:
            st.markdown("#### 🎨 Appearance")
            st.selectbox("Default Theme", ["Light", "Dark", "System"], key="default_theme")
            st.selectbox("Language", ["English", "Spanish", "French"], key="language")
        
        with col_set2:
            st.markdown("#### 📧 Notifications")
            st.checkbox("Email notifications for new features", key="email_notif")
            st.checkbox("Weekly summary of plans", key="weekly_summary")
    
    # Footer - Added to all pages
    # Get social links from settings
    if USE_MODULAR_STRUCTURE and Settings:
        linkedin_url = Settings.LINKEDIN_URL
        github_url = Settings.GITHUB_URL
        email = Settings.EMAIL
        portfolio_url = Settings.PORTFOLIO_URL
    else:
        linkedin_url = "https://www.linkedin.com/in/your-profile"
        github_url = "https://github.com/your-username"
        email = "your.email@example.com"
        portfolio_url = "https://your-portfolio-website.com"
    
    st.markdown("---")
    st.markdown(
        f"""
        <div class="footer">
            <div class="footer-content">
                <div class="footer-section">
                    <h4>📚 AI Lesson Planner</h4>
                    <p>Create comprehensive, engaging lesson plans with AI-powered assistance. Perfect for educators at all levels.</p>
                </div>
                <div class="footer-section">
                    <h4>🔗 Quick Links</h4>
                    <p><a href="#home" style="color: {accent_color};">🏠 Home</a></p>
                    <p><a href="#create-plan" style="color: {accent_color};">📝 Create Plan</a></p>
                    <p><a href="#my-plans" style="color: {accent_color};">📚 My Plans</a></p>
                    <p><a href="#settings" style="color: {accent_color};">⚙️ Settings</a></p>
                </div>
                <div class="footer-section">
                    <h4>💡 Features</h4>
                    <p>🤖 AI-Powered Generation</p>
                    <p>📚 Notes & Quiz Creator</p>
                    <p>📥 Multiple Export Formats</p>
                    <p>💾 Cloud Storage</p>
                </div>
                <div class="footer-section">
                    <h4>📞 Support & Connect</h4>
                    <p><a href="{linkedin_url}" target="_blank" style="color: {accent_color}; text-decoration: none;">💼 LinkedIn</a></p>
                    <p><a href="{github_url}" target="_blank" style="color: {accent_color}; text-decoration: none;">🐙 GitHub</a></p>
                    <p><a href="mailto:{email}" style="color: {accent_color}; text-decoration: none;">📧 Email</a></p>
                    <p><a href="{portfolio_url}" target="_blank" style="color: {accent_color}; text-decoration: none;">🌐 Portfolio</a></p>
                </div>
            </div>
            <div class="footer-bottom">
                <p>© 2024 AI Lesson Planner. All rights reserved. | Made with ❤️ for educators worldwide</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )
