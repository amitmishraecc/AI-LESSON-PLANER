"""
Export utilities for lesson plans
"""
import io
import re

# Export libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import RGBColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    RGBColor = None


def generate_pdf(plan_data):
    """Generate PDF from lesson plan"""
    if not REPORTLAB_AVAILABLE:
        return None
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 139),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    title = Paragraph(f"{plan_data['subject']} - {plan_data['topic']}", title_style)
    story.append(title)
    story.append(Spacer(1, 0.2*inch))
    
    # Metadata
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=11,
        textColor=RGBColor(100, 100, 100),
        alignment=TA_CENTER
    )
    meta_text = f"<b>Grade/Level:</b> {plan_data['grade']} | <b>Duration:</b> {plan_data['duration']} | <b>Created:</b> {plan_data.get('created_at', 'N/A')}"
    story.append(Paragraph(meta_text, meta_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Content - Simple markdown to text conversion
    content = plan_data['content']
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('# '):
                style = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=18, spaceAfter=12, spaceBefore=12)
                story.append(Paragraph(line[2:].strip(), style))
            elif line.startswith('## '):
                style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, spaceAfter=10, spaceBefore=10)
                story.append(Paragraph(line[3:].strip(), style))
            elif line.startswith('### '):
                style = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=12, spaceAfter=8, spaceBefore=8)
                story.append(Paragraph(line[4:].strip(), style))
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f"• {line[2:].strip()}", styles['Normal']))
            elif line.startswith('**') and line.endswith('**'):
                bold_text = line.strip('*')
                story.append(Paragraph(f"<b>{bold_text}</b>", styles['Normal']))
            else:
                # Clean up markdown links
                line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_word_doc(plan_data):
    """Generate Word document from lesson plan"""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f"{plan_data['subject']} - {plan_data['topic']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Grade/Level: {plan_data['grade']}").bold = True
    meta_para.add_run(f" | Duration: {plan_data['duration']}").bold = True
    meta_para.add_run(f" | Created: {plan_data.get('created_at', 'N/A')}").bold = True
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Content
    content = plan_data['content']
    lines = content.split('\n')
    
    for line in lines:
        if line.strip():
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif line.startswith('**') and line.endswith('**'):
                para = doc.add_paragraph()
                para.add_run(line.strip('*')).bold = True
            else:
                # Clean markdown links
                clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                doc.add_paragraph(clean_line)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

